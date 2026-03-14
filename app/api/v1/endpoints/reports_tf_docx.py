"""
app/api/v1/endpoints/reports_tf_docx.py
Route : GET /api/v1/reports/technico-fonctionnel-docx
Génère le rapport technico-fonctionnel au format Word (.docx).
Dépendance : python-docx
"""

import base64
import io
import re
import uuid
from datetime import date, datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import Optional, List, Tuple

from fastapi import APIRouter, Depends, Query
from fastapi.responses import Response
from sqlalchemy.orm import Session, joinedload

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.db.session import get_db
from app.api.deps import get_current_user
from app.models.user import User
from app.models.task import Task, TaskStatus
from app.core.exceptions import NotFoundException
from app.repositories.user_repository import UserRepository

router = APIRouter(prefix="/reports", tags=["Rapports PDF"])

ASSETS_DIR = Path(__file__).parent.parent.parent.parent / "assets" / "report"

# ── Couleurs SABC ──────────────────────────────────────────────────────────────
SABC_RED    = RGBColor(0xC8, 0x10, 0x2E)
SABC_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
SABC_GREY   = RGBColor(0x7F, 0x8C, 0x8D)
SABC_ORANGE = RGBColor(0xE8, 0x60, 0x0A)


# ══════════════════════════════════════════════════════════════════════════════
# Parser HTML → runs python-docx
# ══════════════════════════════════════════════════════════════════════════════

class _HtmlNode:
    """Représente un nœud inline : texte + style."""
    def __init__(self, text: str, bold=False, italic=False, underline=False):
        self.text = text
        self.bold = bold
        self.italic = italic
        self.underline = underline


class _DocxHtmlParser(HTMLParser):
    """
    Parcourt un fragment HTML et produit des actions sur un Document docx.
    Supporte : p, br, strong/b, em/i, u, h1-h3, ul/ol/li, img(base64).
    """

    def __init__(self):
        super().__init__()
        self._bold      = 0
        self._italic    = 0
        self._underline = 0
        self._buf       = ""
        self._actions: List[dict] = []  # liste d'actions à rejouer
        self._in_li     = False
        self._list_type = None   # "ul" ou "ol"
        self._ol_counter = 0

    def handle_starttag(self, tag, attrs):
        attr_dict = dict(attrs)
        if tag in ("p", "div"):
            self._flush("p")
        elif tag in ("h1",):
            self._flush("p"); self._actions.append({"type": "heading", "level": 1, "text": ""})
        elif tag in ("h2",):
            self._flush("p"); self._actions.append({"type": "heading", "level": 2, "text": ""})
        elif tag in ("h3",):
            self._flush("p"); self._actions.append({"type": "heading", "level": 3, "text": ""})
        elif tag in ("strong", "b"):
            self._bold += 1
        elif tag in ("em", "i"):
            self._italic += 1
        elif tag == "u":
            self._underline += 1
        elif tag == "br":
            self._buf += "\n"
        elif tag == "ul":
            self._flush("p"); self._list_type = "ul"
        elif tag == "ol":
            self._flush("p"); self._list_type = "ol"; self._ol_counter = 0
        elif tag == "li":
            self._in_li = True; self._buf = ""
        elif tag == "img":
            src = attr_dict.get("src", "")
            if src.startswith("data:image"):
                self._actions.append({"type": "image", "src": src, "attrs": attr_dict})

    def handle_endtag(self, tag):
        if tag in ("p", "div"):
            self._flush("p")
        elif tag in ("h1", "h2", "h3"):
            level = int(tag[1])
            # Remplir le dernier heading en attente
            text = self._buf.strip(); self._buf = ""
            for a in reversed(self._actions):
                if a.get("type") == "heading" and not a.get("text"):
                    a["text"] = text; break
        elif tag in ("strong", "b"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("em", "i"):
            self._italic = max(0, self._italic - 1)
        elif tag == "u":
            self._underline = max(0, self._underline - 1)
        elif tag == "li":
            text = self._buf.strip(); self._buf = ""
            if self._list_type == "ul":
                self._actions.append({"type": "bullet", "text": text})
            else:
                self._ol_counter += 1
                self._actions.append({"type": "numbered", "text": text, "num": self._ol_counter})
            self._in_li = False
        elif tag in ("ul", "ol"):
            self._list_type = None

    def handle_data(self, data):
        self._buf += data

    def _flush(self, kind: str):
        text = self._buf.strip(); self._buf = ""
        if text:
            self._actions.append({
                "type": "para",
                "text": text,
                "bold": self._bold > 0,
                "italic": self._italic > 0,
                "underline": self._underline > 0,
            })

    @property
    def actions(self):
        if self._buf.strip():
            self._flush("p")
        return self._actions


def _apply_html_to_doc(doc: Document, html: str):
    """Applique le contenu HTML parsé à un document docx."""
    if not html or not html.strip():
        return

    parser = _DocxHtmlParser()
    parser.feed(html)

    for action in parser.actions:
        t = action.get("type")

        if t == "para":
            p = doc.add_paragraph()
            run = p.add_run(action["text"])
            run.bold      = action.get("bold", False)
            run.italic    = action.get("italic", False)
            run.underline = action.get("underline", False)

        elif t == "heading":
            level = action.get("level", 2)
            text  = action.get("text", "")
            if text:
                doc.add_heading(text, level=level)

        elif t == "bullet":
            p = doc.add_paragraph(action["text"], style="List Bullet")

        elif t == "numbered":
            p = doc.add_paragraph(action["text"], style="List Number")

        elif t == "image":
            src = action["src"]
            try:
                match = re.match(r"data:image/(\w+);base64,(.*)", src, re.DOTALL)
                if match:
                    raw = base64.b64decode(match.group(2))
                    buf = io.BytesIO(raw)
                    doc.add_picture(buf, width=Cm(12))
            except Exception:
                pass


def _html_to_plain(html: str) -> str:
    """Extrait le texte brut d'un fragment HTML."""
    class _S(HTMLParser):
        def __init__(self): super().__init__(); self._p = []
        def handle_data(self, d): self._p.append(d)
        def handle_starttag(self, t, a):
            if t in ("p", "br", "li", "div"): self._p.append("\n")
        @property
        def text(self): return " ".join("".join(self._p).split())
    s = _S(); s.feed(html); return s.text.strip()


# ── Helpers style docx ─────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Colorie le fond d'une cellule de tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _add_horizontal_rule(doc: Document):
    """Ajoute un filet horizontal dans le document."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_footer(doc: Document, user: User):
    """Configure le pied de page avec nom du stagiaire et numéro de page."""
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = fp.add_run(f"{user.full_name}  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = SABC_GREY

    # Numéro de page via champ XML
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run2 = fp.add_run()
    run2.font.size = Pt(8)
    run2.font.color.rgb = SABC_GREY
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)


# ── Générateur principal ───────────────────────────────────────────────────────

def _generate_docx(user: User, tasks: List[Task], date_from: Optional[date], date_to: Optional[date]) -> bytes:
    doc = Document()

    # ── Marges ────────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Pied de page ─────────────────────────────────────────────────────────
    _set_footer(doc, user)

    # ── 1. Page de garde ──────────────────────────────────────────────────────
    # Logo SABC (si disponible)
    logo_sabc = ASSETS_DIR / "logo_sabc.png"
    if logo_sabc.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo_sabc), width=Cm(4))

    logo_enspy = ASSETS_DIR / "logo.png"
    if logo_enspy.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo_enspy), width=Cm(4))

    doc.add_paragraph()

    # Titre
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Rapport Technico-Fonctionnel")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = SABC_DARK

    doc.add_paragraph()

    # Infos stagiaire
    period_from = (date_from or user.internship_start_date).strftime("%d/%m/%Y")
    period_to   = (date_to or date.today()).strftime("%d/%m/%Y")

    info_tbl = doc.add_table(rows=5, cols=2)
    info_tbl.style = "Table Grid"
    labels = ["Stagiaire", "Département", "Encadreur", "Début de stage", "Période du rapport"]
    values = [
        user.full_name,
        user.department,
        user.supervisor_name,
        user.internship_start_date.strftime("%d/%m/%Y"),
        f"{period_from} → {period_to}",
    ]
    for i, (label, value) in enumerate(zip(labels, values)):
        lc = info_tbl.rows[i].cells[0]
        vc = info_tbl.rows[i].cells[1]
        _set_cell_bg(lc, "#F2F3F4")
        lc.text = label
        lc.paragraphs[0].runs[0].bold = True
        lc.paragraphs[0].runs[0].font.color.rgb = SABC_RED
        vc.text = value

    doc.add_paragraph()

    gen_p = doc.add_paragraph()
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen_p.add_run(
        f"Généré le {datetime.now(timezone.utc).strftime('%d/%m/%Y à %H:%M')} UTC"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = SABC_GREY

    doc.add_page_break()

    # ── 2. Table des matières ─────────────────────────────────────────────────
    toc_title = doc.add_heading("Table des matières", level=1)
    toc_title.runs[0].font.color.rgb = SABC_DARK

    # Regroupement par domaine
    domain_map: dict = {}
    for task in tasks:
        domain = task.effective_domain_name
        domain_map.setdefault(domain, []).append(task)

    chap_num = 1
    for domain_name, domain_tasks in domain_map.items():
        p = doc.add_paragraph()
        run = p.add_run(f"{chap_num}. {domain_name}")
        run.bold = True; run.font.size = Pt(11)
        sec_num = 1
        for task in domain_tasks:
            p2 = doc.add_paragraph(style="List Bullet")
            p2.paragraph_format.left_indent = Cm(1)
            run2 = p2.add_run(
                f"{chap_num}.{sec_num}  "
                f"{task.task_date.strftime('%d/%m/%Y')} — {task.title}"
            )
            run2.font.size = Pt(10)
            sec_num += 1
        chap_num += 1

    doc.add_page_break()

    # ── 3. Corps : chapitres par domaine ──────────────────────────────────────
    chap_num = 1
    for domain_name, domain_tasks in domain_map.items():
        # Titre chapitre (Heading 1)
        h1 = doc.add_heading(f"Chapitre {chap_num} — {domain_name}", level=1)
        h1.runs[0].font.color.rgb = SABC_RED

        intro = doc.add_paragraph(
            f"Ce chapitre regroupe les activités réalisées dans le domaine "
            f"{domain_name} au cours de la période couverte par ce rapport."
        )
        intro.paragraph_format.space_after = Pt(10)

        sec_num = 1
        for task in domain_tasks:
            # Titre section (Heading 2)
            h2 = doc.add_heading(
                f"{chap_num}.{sec_num}  {task.title}", level=2
            )
            h2.runs[0].font.color.rgb = SABC_DARK

            # Tableau méta
            try:
                duration = (
                    datetime.combine(task.task_date, task.end_time) -
                    datetime.combine(task.task_date, task.start_time)
                ).seconds // 60
                dur_str = f"{duration // 60}h{duration % 60:02d}"
            except Exception:
                dur_str = ""

            meta_tbl = doc.add_table(rows=2, cols=4)
            meta_tbl.style = "Table Grid"

            headers_row = meta_tbl.rows[0].cells
            data_row    = meta_tbl.rows[1].cells

            header_labels = ["Date", "Horaire", "Catégorie", "Statut"]
            data_values = [
                task.task_date.strftime("%d/%m/%Y"),
                f"{task.start_time.strftime('%H:%M')} — {task.end_time.strftime('%H:%M')}"
                + (f" ({dur_str})" if dur_str else ""),
                task.category.value if hasattr(task.category, "value") else str(task.category),
                task.status.value   if hasattr(task.status, "value")   else str(task.status),
            ]
            for i, (lbl, val) in enumerate(zip(header_labels, data_values)):
                _set_cell_bg(headers_row[i], "#C8102E")
                headers_row[i].text = lbl
                headers_row[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                headers_row[i].paragraphs[0].runs[0].bold = True
                data_row[i].text = val

            doc.add_paragraph()

            # Description (si pas de commentaire)
            comment = task.comments[0] if task.comments else None
            if task.description and not comment:
                desc_head = doc.add_paragraph()
                run = desc_head.add_run("Description")
                run.bold = True; run.font.size = Pt(11)
                _apply_html_to_doc(doc, task.description)

            # Commentaire riche
            if comment and comment.content and comment.content.strip():
                com_head = doc.add_paragraph()
                run = com_head.add_run("Commentaire")
                run.bold = True; run.font.size = Pt(11)
                _apply_html_to_doc(doc, comment.content)
            elif not task.description:
                p_none = doc.add_paragraph()
                run = p_none.add_run("Aucun commentaire renseigné pour cette activité.")
                run.italic = True
                run.font.color.rgb = SABC_GREY

            # Image fichier optionnelle
            for ext in (".png", ".jpg", ".jpeg"):
                img_path = ASSETS_DIR / "tasks" / f"{task.id}{ext}"
                if img_path.exists():
                    doc.add_picture(str(img_path), width=Cm(12))
                    caption = doc.add_paragraph(f"Figure {chap_num}.{sec_num} — {task.title}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.runs[0].italic = True
                    break

            _add_horizontal_rule(doc)
            sec_num += 1

        chap_num += 1
        doc.add_page_break()

    # ── 4. Dernière page ───────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    last_p = doc.add_paragraph()
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = last_p.add_run("WakAgenda")
    run.bold = True; run.font.size = Pt(28)
    run.font.color.rgb = SABC_DARK

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = tagline.add_run(
        "Agenda interactif des stagiaires\n"
        "Boissons du Cameroun — Direction des Systèmes d'Information\n"
        "77, rue Prince Bell — Bali, Douala"
    )
    run2.font.size = Pt(11)
    run2.font.color.rgb = SABC_GREY

    # ── Export ────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get(
    "/technico-fonctionnel-docx",
    summary="Rapport technico-fonctionnel au format Word (.docx)",
    response_class=Response,
    responses={
        200: {
            "content": {
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {}
            },
            "description": "Rapport DOCX généré",
        }
    },
)
def generate_rapport_tf_docx(
    date_from: Optional[date] = Query(None, description="Date de début (YYYY-MM-DD)."),
    date_to:   Optional[date] = Query(None, description="Date de fin (YYYY-MM-DD)."),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Génère le rapport technico-fonctionnel au format Word :
    - Page de garde (logos, infos stagiaire)
    - Table des matières
    - Chapitres par domaine, sections par tâche
    - Commentaires HTML → contenu Word formaté (bold, italic, listes, images)
    - Pied de page : nom + numéro de page
    """
    user_repo = UserRepository(db)
    user = user_repo.get_by_id(current_user.id)
    if not user:
        raise NotFoundException("Utilisateur")

    tasks = (
        db.query(Task)
        .options(
            joinedload(Task.custom_domain),
            joinedload(Task.comments),
        )
        .filter(
            Task.user_id == current_user.id,
            Task.task_date >= (date_from or user.internship_start_date),
            Task.task_date <= (date_to   or date.today()),
        )
        .order_by(Task.task_date.asc(), Task.start_time.asc())
        .limit(10_000)
        .all()
    )

    docx_bytes = _generate_docx(user, tasks, date_from, date_to)

    filename = (
        f"rapport_technico_fonctionnel"
        f"_{current_user.last_name.lower()}"
        f"_{date.today().isoformat()}.docx"
    )
    return Response(
        content=docx_bytes,
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )